"""
FleetInsight AI — FastAPI Backend v2.2
========================================
Endpoints :
  POST /analyze         → Analyse complète + Gemini AI
  POST /report          → Rapport PDF / Word
  GET  /health          → Health check
"""

from __future__ import annotations
import io
import os
from datetime import datetime
from typing import Any, Optional

import numpy as np
import pandas as pd
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel

from fleet_analyzer import FleetAnalyzer

try:
    from gemini_advisor import run_gemini_analysis
    GEMINI_OK = True
except ImportError:
    GEMINI_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Table, TableStyle, HRFlowable)
    from reportlab.lib.units import cm
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    from docx import Document
    from docx.shared import Pt, Cm
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ─────────────────────────────────────────────────────────────────────────────
app = FastAPI(
    title="FleetInsight AI",
    description="Analyse intelligente de parc automobile — Powered by Gemini AI",
    version="2.2.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://miensie.github.io","http://localhost:3000","https://localhost:3000","*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─────────────────────────────────────────────────────────────────────────────
class FleetPayload(BaseModel):
    sheets: dict[str, list[dict]]
    use_gemini: bool = True  # activer/désactiver Gemini

class ReportPayload(BaseModel):
    result: dict
    gemini: Optional[dict] = None
    format: str = "pdf"
    title: str = "FleetInsight AI — Rapport de Suivi du Parc Automobile"

# ─────────────────────────────────────────────────────────────────────────────

def _safe(obj: Any) -> Any:
    if isinstance(obj, (np.integer,)): return int(obj)
    if isinstance(obj, (np.floating,)): return None if np.isnan(obj) else float(obj)
    if isinstance(obj, np.ndarray): return obj.tolist()
    if isinstance(obj, dict): return {k: _safe(v) for k, v in obj.items()}
    if isinstance(obj, list): return [_safe(i) for i in obj]
    return obj

def _fmt(n) -> str:
    if n is None or (isinstance(n, float) and (np.isnan(n) or np.isinf(n))): return "—"
    n = float(n)
    if n >= 1_000_000: return f"{n/1_000_000:.1f}M FCFA"
    if n >= 1_000:     return f"{n/1_000:.0f}K FCFA"
    return f"{n:.0f} FCFA"


# ─────────────────────────────────────────────────────────────────────────────
# ENDPOINTS
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {
        "status": "ok",
        "version": "2.2.0",
        "pdf": PDF_OK,
        "docx": DOCX_OK,
        "gemini": GEMINI_OK and bool(os.getenv("GEMINI_API_KEY")),
        "timestamp": datetime.utcnow().isoformat(),
    }


@app.post("/analyze")
async def analyze(payload: FleetPayload):
    if not payload.sheets:
        raise HTTPException(400, "Le classeur est vide.")
    try:
        # Analyse des données Excel
        analyzer = FleetAnalyzer(payload.sheets)
        result   = analyzer.analyze()

        # Analyse Gemini AI (si disponible et activée)
        gemini_result = None
        if payload.use_gemini and GEMINI_OK and os.getenv("GEMINI_API_KEY"):
            try:
                gemini_result = await run_gemini_analysis(result)
                # Mettre à jour l'économie estimée dans l'optimisation
                savings = gemini_result.get("budget_et_economies", {})
                if "economie_estimee" in savings:
                    result["niveau_3_decisions"]["optimisation"]["economie_estimee"] = \
                        savings.get("economie_estimee")
                    result["niveau_3_decisions"]["optimisation"]["economie_fmt"] = \
                        savings.get("economie_fmt", "—")
            except Exception as e:
                gemini_result = {"error": str(e), "gemini_unavailable": True}

        response = {**result, "gemini": gemini_result}
        return JSONResponse(content=_safe(response))
    except Exception as e:
        raise HTTPException(500, f"Erreur d'analyse : {e}")


@app.post("/report")
async def report(payload: ReportPayload):
    fmt = payload.format.lower()
    if fmt == "pdf":
        if not PDF_OK:
            raise HTTPException(501, "ReportLab non installé.")
        content  = _make_pdf(payload)
        media    = "application/pdf"
        filename = "fleetinsight_rapport.pdf"
    elif fmt == "word":
        if not DOCX_OK:
            raise HTTPException(501, "python-docx non installé.")
        content  = _make_word(payload)
        media    = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = "fleetinsight_rapport.docx"
    else:
        raise HTTPException(400, "Format invalide ('pdf' ou 'word')")

    return StreamingResponse(
        io.BytesIO(content), media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─────────────────────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────────────────────

def _make_pdf(payload: ReportPayload) -> bytes:
    buf  = io.BytesIO()
    doc  = SimpleDocTemplate(buf, pagesize=A4,
                              topMargin=2*cm, bottomMargin=2*cm,
                              leftMargin=2*cm, rightMargin=2*cm)
    styles = getSampleStyleSheet()
    C_NAVY = colors.HexColor("#0e1628")
    C_AMB  = colors.HexColor("#f59e0b")
    C_BLUE = colors.HexColor("#1e3354")
    C_RED  = colors.HexColor("#ef4444")
    C_GRN  = colors.HexColor("#10b981")
    C_GRAY = colors.HexColor("#6b7280")

    h1 = ParagraphStyle("H1", parent=styles["Heading1"],
                         textColor=C_BLUE, fontSize=16, spaceAfter=6)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"],
                         textColor=C_BLUE, fontSize=12, spaceAfter=4)
    body = ParagraphStyle("Body", parent=styles["Normal"],
                           fontSize=9, leading=14,
                           textColor=colors.HexColor("#374151"))
    story = []

    r   = payload.result
    ops = r.get("niveau_1_operationnel", {})
    ind = r.get("niveau_2_indicateurs", {})
    dec = r.get("niveau_3_decisions", {})
    gem = payload.gemini or {}
    kpis= r.get("kpis", [])

    # ── Titre ──
    story.append(Paragraph(payload.title, h1))
    story.append(Paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} — Powered by Gemini AI",
                            ParagraphStyle("sub", parent=styles["Normal"], fontSize=8, textColor=C_GRAY)))
    story.append(HRFlowable(width="100%", thickness=2, color=C_AMB, spaceAfter=12))

    # ── Résumé exécutif Gemini ──
    exec_sum = gem.get("resume_executif", {})
    if exec_sum and not exec_sum.get("parse_error"):
        story.append(Paragraph("Résumé Exécutif — Analyse IA", h2))
        if exec_sum.get("message_cle"):
            story.append(Paragraph(f"<b>📌 {exec_sum['message_cle']}</b>",
                                   ParagraphStyle("key", parent=body, textColor=C_AMB)))
        if exec_sum.get("resume_narratif"):
            story.append(Paragraph(exec_sum["resume_narratif"], body))
        story.append(Spacer(1, 8))

    # ── KPIs ──
    story.append(Paragraph("Tableau de Bord Stratégique", h2))
    kpi_data = [["Indicateur", "Valeur", "Niv."]] + [
        [k["label"], k["value"], f"N{k['niveau']}"] for k in kpis
    ]
    kt = Table(kpi_data, colWidths=[8*cm, 5*cm, 3*cm])
    kt.setStyle(TableStyle([
        ("BACKGROUND", (0,0),(-1,0), C_BLUE),
        ("TEXTCOLOR",  (0,0),(-1,0), colors.white),
        ("FONTSIZE",   (0,0),(-1,-1), 9),
        ("ROWBACKGROUNDS",(0,1),(-1,-1), [colors.HexColor("#f9fafb"), colors.white]),
        ("GRID",       (0,0),(-1,-1), 0.5, colors.HexColor("#e5e7eb")),
        ("ALIGN",      (1,0),(-1,-1), "CENTER"),
    ]))
    story.append(kt)
    story.append(Spacer(1, 10))

    # ── Alertes ──
    alertes = dec.get("alertes", [])
    if alertes:
        story.append(Paragraph("Alertes Prioritaires", h2))
        for a in alertes:
            col = C_RED if a["niveau"]=="CRITIQUE" else colors.HexColor("#f59e0b")
            story.append(Paragraph(f"<b>{a['icon']} [{a['domaine']}] {a['titre']}</b>",
                                   ParagraphStyle("al", parent=body, textColor=col, spaceAfter=2)))
            story.append(Paragraph(f"→ {a['action']}", body))
            story.append(Spacer(1, 3))

    # ── VT ──
    vt_op = ops.get("vt", {})
    story.append(Paragraph("Conformité Visites Techniques (VT)", h2))
    story.append(Paragraph(
        f"✅ OUI : {vt_op.get('oui',0)}  |  ❌ NON : {vt_op.get('non',0)}  |  ⏳ PAS ENCORE : {vt_op.get('pas_encore',0)}", body
    ))
    vt_gem = gem.get("conformite_vt", {})
    if vt_gem and not vt_gem.get("parse_error"):
        if vt_gem.get("message_direction"):
            story.append(Paragraph(f"🤖 Gemini : <i>{vt_gem['message_direction']}</i>",
                                   ParagraphStyle("g", parent=body, textColor=C_AMB)))
    story.append(Spacer(1, 8))

    # ── Budget ──
    optim = dec.get("optimisation", {})
    carb  = ops.get("carburant", {})
    entr  = ops.get("entretien", {})
    gen   = ops.get("generateur", {})
    story.append(Paragraph("Budget du Parc", h2))
    bdata = [
        ["Poste","Montant","Part (%)"],
        ["Carburant",       _fmt(carb.get("total_depense")), f"{optim.get('part_carburant_pct',0):.1f}%"],
        ["Entretien",       _fmt(entr.get("total_depense")), f"{optim.get('part_entretien_pct',0):.1f}%"],
        ["Groupe élect.",   _fmt(gen.get("total_depense",0)), f"{optim.get('part_generateur_pct',0):.1f}%"],
        ["TOTAL",           optim.get("total_fmt","—"), "100%"],
    ]
    bt = Table(bdata, colWidths=[7*cm, 6*cm, 3*cm])
    bt.setStyle(TableStyle([
        ("BACKGROUND", (0,0),(-1,0), C_BLUE),
        ("TEXTCOLOR",  (0,0),(-1,0), colors.white),
        ("BACKGROUND", (0,-1),(-1,-1), colors.HexColor("#f0f9ff")),
        ("FONTSIZE",   (0,0),(-1,-1), 9),
        ("GRID",       (0,0),(-1,-1), 0.5, colors.HexColor("#e5e7eb")),
        ("ALIGN",      (1,0),(-1,-1), "RIGHT"),
    ]))
    story.append(bt)

    # ── Économie Gemini ──
    savings = gem.get("budget_et_economies", {})
    if savings and not savings.get("parse_error") and savings.get("economie_fmt"):
        story.append(Spacer(1,6))
        story.append(Paragraph(
            f"🤖 Économie potentielle estimée par Gemini : <b>{savings.get('economie_fmt','—')}</b> "
            f"({savings.get('economie_pct','—')}% du budget)",
            ParagraphStyle("sav", parent=body, textColor=C_GRN)
        ))
        story.append(Paragraph(f"Score de gestion : {savings.get('score_gestion','—')}/100 — {savings.get('niveau_maturite','—')}",body))
        if savings.get("synthese"):
            story.append(Paragraph(savings["synthese"], body))

    # ── Plan renouvellement ──
    plan = dec.get("plan_renouvellement", {})
    story.append(Spacer(1,10))
    story.append(Paragraph("Plan de Renouvellement", h2))
    story.append(Paragraph(
        f"Véhicules à remplacer : {plan.get('nb_remplacement',0)} | "
        f"Budget : {plan.get('budget_fmt','—')} | Horizon : {plan.get('horizon','—')}", body
    ))

    # ── Recommandations ──
    leviers = savings.get("leviers",[]) if savings else []
    if leviers:
        story.append(Spacer(1,10))
        story.append(Paragraph("Leviers d'Optimisation — Gemini AI", h2))
        for l in leviers[:5]:
            story.append(Paragraph(
                f"• [{l.get('priorite','—')}] {l.get('levier','—')} — {_fmt(l.get('economie_potentielle',0))}",
                body
            ))
            story.append(Paragraph(f"  {l.get('detail','')}", ParagraphStyle("det",parent=body,textColor=C_GRAY)))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# WORD
# ─────────────────────────────────────────────────────────────────────────────

def _make_word(payload: ReportPayload) -> bytes:
    doc = Document()
    doc.add_heading(payload.title, 0)
    doc.add_paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} — FleetInsight AI")

    r   = payload.result
    ops = r.get("niveau_1_operationnel", {})
    dec = r.get("niveau_3_decisions", {})
    gem = payload.gemini or {}

    exec_sum = gem.get("resume_executif", {})
    if exec_sum and exec_sum.get("resume_narratif"):
        doc.add_heading("Résumé Exécutif", level=1)
        doc.add_paragraph(exec_sum["resume_narratif"])

    doc.add_heading("Alertes", level=1)
    for a in dec.get("alertes",[]):
        doc.add_paragraph(f"{a['icon']} [{a['domaine']}] {a['titre']}", style="List Bullet")
        doc.add_paragraph(f"   → {a['action']}")

    vt_op = ops.get("vt",{})
    doc.add_heading("Conformité VT", level=1)
    doc.add_paragraph(f"OUI : {vt_op.get('oui',0)}  |  NON : {vt_op.get('non',0)}  |  PAS ENCORE : {vt_op.get('pas_encore',0)}")

    doc.add_heading("Budget", level=1)
    optim = dec.get("optimisation",{})
    doc.add_paragraph(f"Carburant : {_fmt(ops.get('carburant',{}).get('total_depense'))}")
    doc.add_paragraph(f"Entretien : {_fmt(ops.get('entretien',{}).get('total_depense'))}")
    doc.add_paragraph(f"Groupe élect. : {_fmt(ops.get('generateur',{}).get('total_depense',0))}")
    doc.add_paragraph(f"Total : {optim.get('total_fmt','—')}")

    savings = gem.get("budget_et_economies",{})
    if savings and savings.get("economie_fmt"):
        doc.add_heading("Économie Potentielle (Gemini AI)", level=1)
        doc.add_paragraph(f"Économie estimée : {savings.get('economie_fmt','—')}")
        doc.add_paragraph(f"Score de gestion : {savings.get('score_gestion','—')}/100")
        if savings.get("synthese"):
            doc.add_paragraph(savings["synthese"])

    plan = dec.get("plan_renouvellement",{})
    doc.add_heading("Plan de Renouvellement", level=1)
    doc.add_paragraph(f"Véhicules à remplacer : {plan.get('nb_remplacement',0)}")
    doc.add_paragraph(f"Budget estimé : {plan.get('budget_fmt','—')}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()




# ─────────────────────────────────────────────────────────────────────────────
# ENDPOINT : /chat  (Chatbot IA en temps réel)
# ─────────────────────────────────────────────────────────────────────────────

class ChatPayload(BaseModel):
    question: str
    analysis: dict   # résultat /analyze stocké côté client

@app.post("/chat", tags=["Chat"])
async def chat(payload: ChatPayload):
    """
    Répond à une question en langage naturel sur le parc auto via Gemini AI.
    Nécessite le résultat d'analyse passé en contexte.
    """
    if not payload.question.strip():
        raise HTTPException(400, "La question est vide.")

    if not GEMINI_OK or not os.getenv("GEMINI_API_KEY"):
        # Fallback local si IA non dispo
        return JSONResponse(content={
            "reponse": _local_fallback(payload.question, payload.analysis),
            "ia_used": False,
        })

    try:
        from gemini_advisor import chat_with_ia
        result = await chat_with_ia(payload.question, payload.analysis)
        return JSONResponse(content={
            "reponse": result.get("reponse", "Pas de réponse."),
            "ia_used": not result.get("error", False),
        })
    except Exception as e:
        return JSONResponse(content={
            "reponse": _local_fallback(payload.question, payload.analysis),
            "ia_used": False,
            "error": str(e),
        })


def _local_fallback(question: str, analysis: dict) -> str:
    """Réponses locales si Gemini indisponible — utilise les données réelles."""
    q   = question.lower()
    ops = analysis.get("niveau_1_operationnel", {})
    ind = analysis.get("niveau_2_indicateurs",  {})
    dec = analysis.get("niveau_3_decisions",     {})
    vt  = ops.get("vt", {})
    conf= ind.get("conformite_vt", {})
    taux= ind.get("taux_immobilisation", {})
    carb= ops.get("carburant", {})
    entr= ops.get("entretien", {})
    gen = ops.get("generateur", {})
    optim = dec.get("optimisation", {})

    if any(k in q for k in ["vt","visite","conformité","conformite"]):
        return (f"Conformité VT : {conf.get('taux_conformite','—')}% — "
                f"{vt.get('oui',0)} OK, {vt.get('non',0)} non faites, "
                f"{vt.get('pas_encore',0)} pas encore. "
                f"Risque : {conf.get('niveau_risque','—')}.")
    if any(k in q for k in ["budget","coût","cout","dépense","economie","économie"]):
        return (f"Budget total parc : {optim.get('total_fmt','—')}. "
                f"Carburant : {_fmt(carb.get('total_depense'))}, "
                f"Entretien : {_fmt(entr.get('total_depense'))}, "
                f"Groupe élect. : {_fmt(gen.get('total_depense',0))}.")
    if any(k in q for k in ["atelier","immobilisation","disponibilité","disponibilite"]):
        return (f"Disponibilité du parc : {taux.get('taux_dispo','—')}%. "
                f"{taux.get('en_atelier',0)} véhicule(s) actuellement en atelier.")
    if any(k in q for k in ["alerte","urgent","critique","priorité","priorite"]):
        alertes = dec.get("alertes", [])
        crit = [a for a in alertes if a.get("niveau") == "CRITIQUE"]
        if crit:
            return f"{len(crit)} alerte(s) critique(s) : " + " | ".join(a.get("titre","") for a in crit[:3])
        return f"{len(alertes)} alerte(s) active(s) — aucune critique."
    if any(k in q for k in ["électrogène","electrogene","generateur","générateur"]):
        return (f"Groupe électrogène : {_fmt(gen.get('total_depense',0))} de dépenses, "
                f"{gen.get('nb_entrees',0)} transaction(s)." if gen.get("disponible")
                else "Données groupe électrogène non disponibles.")
    if any(k in q for k in ["carburant","essence","gasoil","fuel"]):
        top = carb.get("top_consommateurs", [])
        top1 = top[0] if top else None
        return (f"Budget carburant : {_fmt(carb.get('total_depense'))} "
                f"({carb.get('nb_transactions',0)} transactions). "
                + (f"Top consommateur : {top1['vehicule']} ({top1['total_fmt']})." if top1 else ""))
    return ("Je n'ai pas pu analyser cette question précisément. "
            "Consultez les onglets Décisions ou Tableau de bord pour plus de détails.")



# ─────────────────────────────────────────────────────────────────────────────
# ENDPOINT : /vehicle  (Fiche véhicule par immatriculation)
# ─────────────────────────────────────────────────────────────────────────────

class VehiclePayload(BaseModel):
    sheets: dict[str, list[dict]]
    immat: str

@app.post("/vehicle", tags=["Vehicle"])
async def vehicle_search(payload: VehiclePayload):
    """
    Recherche toutes les informations d'un véhicule par immatriculation.
    Croise : LISTE DE VEH, SUIVI_CARBURANT, ENTRETIEN, VT, SORTIES VEH.
    """
    if not payload.immat.strip():
        raise HTTPException(400, "Immatriculation vide.")
    try:
        from vehicle_search import search_vehicle
        result = search_vehicle(payload.sheets, payload.immat)
        return JSONResponse(content=result)
    except Exception as e:
        raise HTTPException(500, f"Erreur de recherche : {e}")


@app.get("/vehicle/list", tags=["Vehicle"])
async def vehicle_list():
    """Retourne la liste de toutes les immatriculations connues (pour autocomplete)."""
    # Cette endpoint est appelée avec les sheets en POST
    return JSONResponse(content={"message": "Utilisez POST /vehicle avec la liste des feuilles."})

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
